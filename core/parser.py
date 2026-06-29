"""
多格式文档解析器 — RAGFlow 离线时的本地降级方案

支持格式：PDF / DOCX / TXT / MD / JPG / PNG
"""
import base64
from pathlib import Path
from typing import Optional


def parse_file(file_path: str) -> tuple[str, dict]:
    """解析文档，返回 (纯文本, 元数据)

    Args:
        file_path: 文件路径

    Returns:
        (text, metadata) — text 为提取的纯文本，metadata 含 file_type/page_count 等
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if not path.exists():
        return "", {"error": "文件不存在: %s" % file_path}

    metadata = {
        "file_name": path.name,
        "file_type": suffix.lstrip("."),
        "file_size": path.stat().st_size,
    }

    if suffix == ".pdf":
        text, meta = _parse_pdf(file_path)
        metadata.update(meta)
        return text, metadata

    if suffix in (".docx", ".doc"):
        text, meta = _parse_docx(file_path)
        metadata.update(meta)
        return text, metadata

    if suffix in (".txt", ".md", ".csv"):
        text = path.read_text(encoding="utf-8", errors="replace")
        metadata["line_count"] = text.count("\n") + 1
        return text, metadata

    if suffix in (".jpg", ".jpeg", ".png", ".bmp", ".webp"):
        text = _parse_image(file_path)
        return text, metadata

    return "", {"error": "不支持的文件格式: %s" % suffix}


# ═══════════════════════════════════════════════════════════
# PDF
# ═══════════════════════════════════════════════════════════

def _parse_pdf(file_path: str) -> tuple[str, dict]:
    """PyMuPDF 提取 PDF 文本"""
    try:
        import fitz
        doc = fitz.open(file_path)
        pages = []
        for page in doc:
            text = page.get_text().strip()
            if text:
                pages.append(text)
        doc.close()
        return "\n\n".join(pages), {"page_count": len(doc)}
    except Exception as e:
        return "[PDF解析失败: %s]" % e, {"error": str(e)}


# ═══════════════════════════════════════════════════════════
# DOCX
# ═══════════════════════════════════════════════════════════

def _parse_docx(file_path: str) -> tuple[str, dict]:
    """python-docx 提取 Word 文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        # 段落文本
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())
        # 表格文本
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs), {"paragraph_count": len(paragraphs)}
    except Exception as e:
        return "[DOCX解析失败: %s]" % e, {"error": str(e)}


# ═══════════════════════════════════════════════════════════
# 图片 OCR
# ═══════════════════════════════════════════════════════════

def _parse_image(file_path: str) -> str:
    """DashScope Qwen-VL 视觉模型 OCR"""
    try:
        from openai import OpenAI
        from config import DASHSCOPE_API_KEY, DASHSCOPE_BASE_URL, QWEN_VL_MODEL

        with open(file_path, "rb") as f:
            img_b64 = base64.b64encode(f.read()).decode()

        suffix = Path(file_path).suffix.lstrip(".").lower()
        mime_type = "image/%s" % ("jpeg" if suffix == "jpg" else suffix)

        client = OpenAI(api_key=DASHSCOPE_API_KEY, base_url=DASHSCOPE_BASE_URL)
        resp = client.chat.completions.create(
            model=QWEN_VL_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{img_b64}"}
                    },
                    {
                        "type": "text",
                        "text": "请提取并输出这张图片中的所有文字内容，保持原有格式和结构。只输出文字，不要额外说明。"
                    },
                ],
            }],
            max_tokens=4096,
        )
        return resp.choices[0].message.content or ""
    except Exception as e:
        return "[OCR错误: %s]" % e
